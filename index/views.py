from django.shortcuts import render
from django.contrib import messages
import random
from django.shortcuts import render
from .models import Ticket
from django.http import FileResponse, HttpResponseNotFound
from io import BytesIO
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import pythoncom
from win32com.client import Dispatch

def replace_text_in_word_file(input_file, output_file, old_text1, new_text1,old_text2,new_text2):
    doc = Document(input_file)

    for paragraph in doc.paragraphs:
        if old_text1 in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text1, "")
            run = paragraph.add_run(new_text1)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        if old_text2 in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text2, "")
            run = paragraph.add_run(new_text2)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text1 in cell.text:
                    cell.text = cell.text.replace(old_text1, "")
                    run = cell.paragraphs[0].add_run(new_text1)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                
                if old_text2 in cell.text:
                    cell.text = cell.text.replace(old_text2, "")
                    run = cell.paragraphs[0].add_run(new_text2)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)

    doc.save(output_file)
    convert_to_pdf(output_file)


def convert_to_pdf(docx_file):
    pythoncom.CoInitialize()
    word = Dispatch('Word.Application')
    docx_file = os.path.abspath(docx_file)  # Convert to absolute path
    doc = word.Documents.Open(docx_file)
    pdf_file = os.path.splitext(docx_file)[0] + '.pdf'
    pdf_file = os.path.abspath(pdf_file)  # Convert to absolute path
    doc.SaveAs(pdf_file, FileFormat=17)  # 17 is the PDF file format
    doc.Close()
    word.Quit()


def save(request):
    return render(request, 'index.html')


def process_form(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        surname = request.POST.get('surname')
        random_number = request.POST.get('randomNumber')

        # Check if a user with the same name and surname already exists
        existing_users = Ticket.objects.filter(name=name, surname=surname)
        if existing_users.exists():
            error_message = "A user with the same name and surname already exists."
            context = {'error_message': error_message}
        else:
            # Save the new user or perform any other required actions
            user = Ticket(name=name, surname=surname, ticket_number=random_number)
            user.save()

            # Prepare context for successful submission
            context = {'name': name, 'surname': surname, 'random_number': random_number}
        

        input_file_path = 'input.docx'
        output_file_path = 'output.docx'

        old_text1 = 'Participant : Turyan Azizov'
        new_text1 = f'Participant : {name} {surname}'

        old_text2 = 'Ticket number : 123454321'
        new_text2 = f'Ticket number : {random_number}'

        replace_text_in_word_file(input_file_path, output_file_path, old_text1, new_text1,old_text2,new_text2)

        pdf_path = 'output.pdf'
        if os.path.exists(pdf_path):
            # Read the PDF content into a BytesIO buffer
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Create and return the FileResponse using the buffer
            buffer = BytesIO(pdf_content)
            response = FileResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="ticket.pdf"'
            return response
    return render(request, 'index.html')